import streamlit as st
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from io import BytesIO
from docx import Document
from openai import AzureOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document as LDocument
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

# Load environment variables
load_dotenv()
KNOWLEDGE_FOLDER = "Knowledge_Repo"


# --- Utility Functions ---

def extract_text(file):
    """Extract text from PDF or DOCX"""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""



def build_knowledge_base(folder=KNOWLEDGE_FOLDER):
    """Build a vector DB from all reference docs in the Knowledge_Repo folder"""
    docs = []
    for f in os.listdir(folder):
        if f.endswith((".pdf", ".docx")):
            path = os.path.join(folder, f)
            text = extract_text(open(path, "rb"))
            docs.append(LDocument(page_content=text, metadata={"source": f}))


    embedding_model = AzureOpenAIEmbeddings(
        model="text-embedding-ada-002",
        azure_endpoint=os.getenv("AZURE_OPENAI_EMD_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_EMD_KEY"),
        api_version=os.getenv("AZURE_OPENAI_EMD_VERSION")
    )

    vectordb = Chroma.from_documents(
        documents=docs,
        embedding=embedding_model,
        collection_name="rfp_responses"
    )
    return vectordb



def create_docx_from_text(text):
    """Generate a clean DOCX from plain text"""
    doc = Document()
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph()
    return doc


def chunk_text(text, max_chars=5000):
    """Split text into roughly equal-sized chunks"""
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 > max_chars:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = para
        else:
            current_chunk += "\n" + para
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks


# --- Summarization (parallelized) ---

def summarize_chunk(chunk, client, i):
    prompt = f"""
Summarize the following RFP section into key requirements, objectives, and scope:
{chunk}
"""
    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}]
    )
    return i, response.choices[0].message.content.strip()


def summarize_rfp_chunks(rfp_text, max_chars=8000, max_workers=4):
    chunks = chunk_text(rfp_text, max_chars=max_chars)
    summaries = [""] * len(chunks)
    progress = st.progress(0)
    status = st.empty()


    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_SUM_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_SUM_KEY"),
        api_version=os.getenv("AZURE_OPENAI_SUM_VERSION")
    )

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(summarize_chunk, chunk, client, i): i for i, chunk in enumerate(chunks)}
        for completed_count, future in enumerate(as_completed(futures), start=1):
            i, summary = future.result()
            summaries[i] = summary
            progress.progress(completed_count / len(chunks))
            status.text(f"Summarizing RFP section {completed_count}/{len(chunks)}...")

    status.text("✅ Summarization complete!")
    return summaries  # Return list of summaries, not joined string


# --- Combine Summaries into One Final Summary ---

def generate_final_summary(chunk_summaries):
    """Combine multiple chunk summaries into one cohesive final RFP summary"""
    # st.info("Creating final condensed RFP summary...")
    
    placeholder = st.empty()  # Create an empty placeholder
    placeholder.info("Creating final condensed RFP summary...")
    time.sleep(4)  # Display for 2 seconds (adjust as needed)
    placeholder.empty()  # Remove the info message

    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FSUM_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FSUM_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FSUM_VERSION")
    )

    combined_text = "\n\n".join(chunk_summaries)

    prompt = f"""
You are an expert in summarizing Request for Proposal (RFP) documents.

Below are individual summaries of different sections of an RFP:
{combined_text}

Please consolidate them into ONE single, cohesive summary that:
- Merges overlapping points
- Highlights scope, objectives, and key requirements
- Removes repetition
- Is concise, clear, and professional (around 1–2 pages)
"""

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.2,
        max_tokens=1200,
        messages=[
            {"role": "system", "content": "You are a professional RFP summarizer."},
            {"role": "user", "content": prompt}
        ]
    )

    final_summary = response.choices[0].message.content.strip()
    st.success("✅ Final condensed summary created!")
    return final_summary




def generate_final_rfp_response(reference_text, condensed_rfp):
    st.info("Generating final RFP response...")

    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )


    prompt = f"""
    You are an expert RFP response writer. 
    Using the Reference Material below, and the RFP Content, generate a **fully structured RFP response**. 

    Rules:
    1. Use # for main sections (like Executive Summary, Scope, Objectives, Functional Requirements, Non-Functional Requirements, Appendix).  
    2. Use ## for subsections.  
    3. Use **•** for bullet points.  
    4. Use **[TABLE]** markers for tables, with proper Markdown table format.
    5. Keep it professional, concise, and clear.  
    6. Include all key points, merging overlapping items and avoiding repetition.

    Reference Material:
    {reference_text}

    RFP Content:
    {condensed_rfp}
    """

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )

    st.success("✅ RFP Response completed!")
    return response.choices[0].message.content.strip()

# --- Insert RFP into Template (Plain Text) ---
def insert_rfp_into_template_plain(template_path, rfp_text, placeholder="<<INSERT_HERE>>"):
    """
    Inserts the RFP response text into a template starting at a placeholder.
    No styling is applied; text is inserted as plain paragraphs.

    Args:
        template_path: Path to the existing Word template.
        rfp_text: Text of the RFP response.
        placeholder: Marker in the template where content should start.
    Returns:
        A python-docx Document object.
    """
    doc = Document(template_path)

    # Find the placeholder paragraph
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            start_index = i
            break
    else:
        # If placeholder not found, append at the end
        start_index = len(doc.paragraphs)
        doc.add_paragraph("")  # ensure there is space

    # Remove the placeholder paragraph
    if start_index < len(doc.paragraphs):
        p = doc.paragraphs[start_index]
        p.clear()

    # Insert RFP response line by line with heading detection
    for line in rfp_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Check for main heading (e.g., "# Executive Summary")
        if stripped.startswith("# "):
            heading_text = stripped.lstrip("# ").strip()
            doc.add_paragraph(heading_text, style="Heading 1")
        elif stripped.startswith("## "):  # optional: subheading
            heading_text = stripped.lstrip("#").strip()
            doc.add_paragraph(heading_text, style="Heading 2")
        else:
            doc.add_paragraph(stripped)

    return doc


import time

#new ui
st.set_page_config(page_title="📄 RFP Response Generator", layout="wide")
st.markdown("<h1 style='text-align:center; color:#4B0082;'>📄 RFP Response Generator</h1>", unsafe_allow_html=True)
st.markdown("---")

st.subheader("Step 1: Upload RFP Document")
uploaded_file = st.file_uploader("Choose PDF or DOCX", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("🔍 Extracting RFP content..."):
        time.sleep(2)
        rfp_text = extract_text(uploaded_file)
    st.success("✅ RFP content extracted!")
    st.text_area("Preview RFP Text", rfp_text[:2000] + "..." if len(rfp_text) > 2000 else rfp_text, height=300)

    # --- Step 2: Build Knowledge Base ---
    st.subheader("Step 2: Build Knowledge Base & Retrieve Reference Material")
    with st.spinner("📚 Building knowledge base..."):
        time.sleep(5)
        knowledge_db = build_knowledge_base()
    st.success("✅ Knowledge base ready!")

    with st.spinner("🔎 Retrieving reference material..."):
        time.sleep(2)
        retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
        ref_docs = retriever.get_relevant_documents(rfp_text)
        reference_text = "\n\n".join([d.page_content for d in ref_docs])
    st.success("✅ Reference material retrieved!")

    # --- Step 3: Summarize RFP ---
    st.subheader("Step 3: Summarize RFP")
    with st.spinner("📝 Summarizing RFP in chunks..."):
        chunk_summaries = summarize_rfp_chunks(rfp_text, max_chars=8000)
    st.success("✅ RFP summarized successfully!")

    # st.subheader("Step 4: Generate Final RFP Response")
    with st.spinner("🔗 Combining summaries into condensed RFP..."):
        condensed_rfp = generate_final_summary(chunk_summaries)
    st.success("✅ Condensed RFP ready!")

    with st.expander("View Chunk Summaries"):
        st.write(condensed_rfp)

    # --- Step 4: Generate Final RFP Response ---
    st.subheader("Step 4: Generate Final RFP Response")

    with st.spinner("✍️ Generating RFP response..."):
        rfp_response = generate_final_rfp_response(reference_text, condensed_rfp)
    st.success("✅ RFP Response generated!")

    st.text_area("Preview RFP Response", rfp_response, height=400)



    # --- Step 5: Download ---
    template_path = "Template/PIPO TO IS Template.docx"
    doc = insert_rfp_into_template_plain(template_path, rfp_response)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 Download RFP Response in Template (DOCX)",
        data=buffer,
        file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

